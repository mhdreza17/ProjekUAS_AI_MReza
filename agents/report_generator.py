import os
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import json
from .base_agent import BaseAgent

class ReportGeneratorAgent(BaseAgent):
    """Enhanced Agent untuk generate laporan audit compliance yang lebih informatif"""
    
    def __init__(self):
        super().__init__("ReportGenerator")
        self.agent_name = "📄 Enhanced Report Generator Agent"
        self.agent_role = "Membuat laporan audit yang komprehensif, actionable, dan mudah dipahami"
        self.reports_dir = "reports"
        os.makedirs(self.reports_dir, exist_ok=True)
        
        print(f"✅ {self.agent_name} initialized")
    
    def process(self, analysis_data: dict, session_id: str):
        """Generate laporan audit profesional dengan enhanced insights"""
        self.set_status("generating")
        self.log_action("Starting enhanced report generation", f"Session: {session_id}")
        
        try:
            # Generate both DOCX and PDF with enhanced content
            docx_path = self.generate_enhanced_docx_report(analysis_data, session_id)
            pdf_path = self.generate_enhanced_pdf_report(analysis_data, session_id)
            
            self.set_status("completed")
            self.log_action("Enhanced report generation completed", f"DOCX: {docx_path}, PDF: {pdf_path}")
            
            return {
                'success': True,
                'docx_path': docx_path,
                'pdf_path': pdf_path,
                'session_id': session_id,
                'analysis_summary': self._create_analysis_summary(analysis_data)
            }
            
        except Exception as e:
            self.set_status("error")
            self.log_action("Report generation error", str(e))
            return {
                'success': False,
                'error': str(e)
            }

    def generate_enhanced_docx_report(self, analysis_data: dict, session_id: str) -> str:
        """Generate enhanced DOCX report dengan struktur yang lebih baik"""
        filename = f"ReguBot_Audit_Report_{session_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.reports_dir, filename)
        
        doc = Document()
        self._setup_document_properties(doc, analysis_data.get('original_filename', 'document'))
        
        # 1. Cover Page
        self._add_enhanced_cover_page(doc, analysis_data, session_id)
        
        # 2. Executive Summary dengan insights
        self._add_enhanced_executive_summary(doc, analysis_data)
        
        # 3. Document Analysis Overview
        self._add_document_analysis_overview(doc, analysis_data)
        
        # 4. Compliance Score Breakdown
        self._add_compliance_score_breakdown(doc, analysis_data)
        
        # 5. Detailed Findings dengan evidence
        self._add_enhanced_detailed_findings(doc, analysis_data)
        
        # 6. Risk Assessment dan Impact Analysis
        self._add_risk_assessment_analysis(doc, analysis_data)
        
        # 7. Actionable Recommendations dengan prioritas
        self._add_prioritized_recommendations(doc, analysis_data)
        
        # 8. Implementation Roadmap
        self._add_implementation_roadmap(doc, analysis_data)
        
        # 9. Appendix dengan referensi lengkap
        self._add_enhanced_appendix(doc, analysis_data)
        
        doc.save(filepath)
        return filepath

    def generate_enhanced_pdf_report(self, analysis_data: dict, session_id: str) -> str:
        """Generate enhanced PDF report dengan visualisasi yang lebih baik"""
        filename = f"compliance_report_{session_id}.pdf"
        filepath = os.path.join(self.reports_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=A4, leftMargin=0.75*inch, rightMargin=0.75*inch)
        styles = getSampleStyleSheet()
        elements = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            spaceAfter=30,
            alignment=1,
            textColor=colors.darkblue
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue,
            borderWidth=1,
            borderColor=colors.lightgrey,
            borderPadding=5
        )
        
        # 1. Cover Page
        elements.extend(self._create_pdf_cover_page(analysis_data, session_id, title_style, styles))
        elements.append(PageBreak())
        
        # 2. Executive Summary
        elements.extend(self._create_pdf_executive_summary(analysis_data, heading_style, styles))
        
        # 3. Compliance Dashboard
        elements.extend(self._create_pdf_compliance_dashboard(analysis_data, heading_style, styles))
        
        # 4. Detailed Analysis
        elements.extend(self._create_pdf_detailed_analysis(analysis_data, heading_style, styles))
        
        # 5. Recommendations
        elements.extend(self._create_pdf_recommendations(analysis_data, heading_style, styles))
        
        # 6. Appendix
        elements.extend(self._create_pdf_appendix(analysis_data, heading_style, styles))
        
        doc.build(elements)
        return filepath

    def _setup_document_properties(self, doc, filename):
        """Enhanced document properties setup"""
        core_props = doc.core_properties
        core_props.title = "ReguBot Enhanced Compliance Audit Report"
        core_props.author = "ReguBot AI Multi-Agent System"
        core_props.subject = f"Comprehensive Compliance Analysis for {filename}"
        core_props.comments = "Generated by ReguBot Enhanced Multi-Agent AI System with Adaptive Analysis"
        core_props.created = datetime.now()
        core_props.keywords = "compliance, audit, GDPR, privacy, data protection, AI analysis"

    def _add_enhanced_appendix(self, doc, analysis_data: dict):
        """Enhanced appendix dengan referensi dan metadata lengkap"""
        doc.add_heading('📚 APPENDIX', level=1)
        
        # Analysis metadata
        doc.add_heading('🔍 Analysis Metadata', level=2)
        metadata_table = doc.add_table(rows=8, cols=2)
        metadata_table.style = 'Light Grid Accent 1'
        
        standards_analyzed = analysis_data.get('analyzed_standards', [])
        doc_analysis = analysis_data.get('document_analysis', {})
        
        metadata_data = [
            ['Standards Analyzed', ', '.join(standards_analyzed) if standards_analyzed else 'General Best Practices'],
            ['Analysis Method', 'AI Multi-Agent Adaptive Analysis'],
            ['System Version', 'ReguBot v2.0 Enhanced'],
            ['Analysis Engine', 'Groq Llama 3.1 + Custom Compliance Logic'],
            ['Document Language', doc_analysis.get('language', 'Unknown')],
            ['Total Aspects Analyzed', str(len(analysis_data.get('detailed_findings', [])))],
            ['Processing Time', 'Real-time AI Analysis'],
            ['Data Security', 'Offline Processing - No Data Transmitted']
        ]
        
        for i, (label, value) in enumerate(metadata_data):
            metadata_table.cell(i, 0).text = label
            metadata_table.cell(i, 1).text = value
            metadata_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        
        # Standards references
        if standards_analyzed:
            doc.add_heading('📋 Standards References', level=2)
            
            standards_info = {
                'GDPR': 'General Data Protection Regulation (EU) 2016/679',
                'UU_PDP': 'Undang-Undang No. 27 Tahun 2022 tentang Perlindungan Data Pribadi',
                'POJK': 'Peraturan Otoritas Jasa Keuangan tentang Perlindungan Konsumen',
                'BSSN': 'Peraturan BSSN tentang Keamanan Siber',
                'NIST': 'NIST Cybersecurity Framework'
            }
            
            for standard in standards_analyzed:
                if standard in standards_info:
                    doc.add_paragraph(f"• {standard}: {standards_info[standard]}")
        
        # Compliance framework explanation
        doc.add_heading('🎯 Compliance Framework', level=2)
        doc.add_paragraph("""
ReguBot menggunakan adaptive compliance framework yang:

1. DOCUMENT ANALYSIS: Menganalisis struktur, jenis, dan konten dokumen
2. RELEVANCE MAPPING: Menentukan aspek compliance yang relevan
3. WEIGHTED SCORING: Memberikan bobot berbeda untuk setiap aspek
4. EVIDENCE EXTRACTION: Mencari bukti konkret dalam dokumen  
5. CONFIDENCE SCORING: Menilai tingkat keyakinan analisis
6. CONTEXTUAL RECOMMENDATIONS: Memberikan rekomendasi yang spesifik

Framework ini memastikan analisis yang:
• Adaptive terhadap jenis dokumen
• Focused pada aspek yang relevan
• Evidence-based dan objektif
• Actionable dan praktis
""")
        
        # Disclaimer
        doc.add_heading('⚖️ Important Disclaimer', level=2)
        doc.add_paragraph("""
DISCLAIMER: Laporan ini dihasilkan oleh sistem AI dan dimaksudkan sebagai alat bantu analisis awal. 
Hasil analisis tidak menggantikan konsultasi dengan ahli hukum atau compliance officer yang qualified. 
Pengguna disarankan untuk:

• Melakukan review manual terhadap temuan AI
• Berkonsultasi dengan legal expert untuk interpretation yang akurat
• Mempertimbangkan konteks bisnis dan jurisdiksi yang spesifik
• Melakukan update berkala seiring perubahan regulasi

ReguBot AI System tidak bertanggung jawab atas keputusan bisnis atau legal yang diambil berdasarkan laporan ini.
""")
        
        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph(f"""
---
Laporan dihasilkan oleh ReguBot Enhanced AI System
Tanggal: {datetime.now().strftime('%d %B %Y, %H:%M:%S WIB')}
Session ID: {analysis_data.get('session_id', 'Unknown')}
© 2024 ReguBot - AI Compliance Checker
---
""")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _create_analysis_summary(self, analysis_data: dict) -> dict:
        """Create comprehensive analysis summary"""
        compliance_score = analysis_data.get('compliance_score', 0)
        
        if compliance_score >= 85:
            compliance_level = "EXCELLENT"
            status_icon = "🟢"
            business_impact = "Low compliance risk. Document demonstrates strong adherence to regulatory requirements."
            immediate_actions = "Continue current practices. Schedule quarterly reviews."
        elif compliance_score >= 70:
            compliance_level = "GOOD" 
            status_icon = "🟡"
            business_impact = "Moderate compliance risk. Some areas need attention but overall framework is solid."
            immediate_actions = "Address identified gaps within 30 days. Enhance documentation."
        elif compliance_score >= 50:
            compliance_level = "NEEDS IMPROVEMENT"
            status_icon = "🟠"
            business_impact = "Significant compliance gaps present. Risk of regulatory issues or user complaints."
            immediate_actions = "Immediate review required. Address high-priority issues within 7 days."
        else:
            compliance_level = "CRITICAL"
            status_icon = "🔴"
            business_impact = "High compliance risk. Significant regulatory and reputational exposure."
            immediate_actions = "Urgent action required. Halt non-compliant processes immediately."
        
        return {
            'compliance_level': compliance_level,
            'status_icon': status_icon,
            'business_impact': business_impact,
            'immediate_actions': immediate_actions
        }

    # PDF-specific methods
    def _create_pdf_cover_page(self, analysis_data: dict, session_id: str, title_style, styles) -> list:
        """Create PDF cover page elements"""
        elements = []
        
        elements.append(Paragraph("🛡️ LAPORAN AUDIT COMPLIANCE", title_style))
        elements.append(Paragraph("ReguBot Enhanced Multi-Agent AI Analysis", styles['Heading2']))
        elements.append(Spacer(1, 12))
        
        # Info table
        info_data = [
            ['📄 Dokumen', analysis_data.get('original_filename', 'document')],
            ['📊 Compliance Score', f"{analysis_data.get('compliance_score', 0)}%"],
            ['🎯 Issues', f"{len(analysis_data.get('issues', []))}"],
            ['✅ Compliant Items', f"{len(analysis_data.get('compliant_items', []))}"],
            ['📅 Tanggal', datetime.now().strftime('%d %B %Y')],
            ['⏰ Waktu', datetime.now().strftime('%H:%M:%S WIB')],
            ['🤖 AI Engine', 'ReguBot v2.0 Enhanced'],
            ['🔒 Security', 'Offline Processing']
        ]
        
        table = Table(info_data, colWidths=[2*inch, 3*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.whitesmoke),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        elements.append(table)
        elements.append(Spacer(1, 24))
        
        return elements

    def _create_pdf_executive_summary(self, analysis_data: dict, heading_style, styles) -> list:
        """Create PDF executive summary elements"""
        elements = []
        
        elements.append(Paragraph('📋 EXECUTIVE SUMMARY', heading_style))
        
        analysis_summary = self._create_analysis_summary(analysis_data)
        doc_analysis = analysis_data.get('document_analysis', {})
        
        summary_text = f"""
{analysis_summary['status_icon']} COMPLIANCE STATUS: {analysis_summary['compliance_level']}

Analisis komprehensif menggunakan ReguBot Enhanced AI telah dilakukan dengan hasil:
• Compliance Score: {analysis_data.get('compliance_score', 0)}% 
• Aspek dianalisis: {len(analysis_data.get('detailed_findings', []))} (adaptive)
• Document type: {doc_analysis.get('document_type', 'Unknown')}
• Main themes: {', '.join(doc_analysis.get('themes', [])[:3])}

{analysis_summary['business_impact']}

IMMEDIATE ACTIONS:
{analysis_summary['immediate_actions']}
        """
        
        elements.append(Paragraph(summary_text.strip(), styles['Normal']))
        elements.append(Spacer(1, 12))
        
        return elements

    def _create_pdf_compliance_dashboard(self, analysis_data: dict, heading_style, styles) -> list:
        """Create PDF compliance dashboard elements"""
        elements = []
        
        elements.append(Paragraph('📊 COMPLIANCE DASHBOARD', heading_style))
        
        # Main metrics
        dashboard_data = [
            ['Metric', 'Value', 'Status'],
            ['Overall Score', f"{analysis_data.get('compliance_score', 0)}%", '🎯'],
            ['Total Aspects', str(len(analysis_data.get('detailed_findings', []))), '🔍'],
            ['Compliant Items', str(len(analysis_data.get('compliant_items', []))), '✅'],
            ['Non-Compliant', str(len(analysis_data.get('issues', []))), '❌'],
            ['High Risk Issues', str(len([i for i in analysis_data.get('issues', []) if i.get('severity') == 'HIGH'])), '🚨']
        ]
        
        dashboard_table = Table(dashboard_data)
        dashboard_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.whitesmoke),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        
        elements.append(dashboard_table)
        elements.append(Spacer(1, 12))
        
        return elements

    def _create_pdf_detailed_analysis(self, analysis_data: dict, heading_style, styles) -> list:
        """Create PDF detailed analysis elements"""
        elements = []
        
        elements.append(Paragraph('🔍 DETAILED FINDINGS', heading_style))
        
        for i, finding in enumerate(analysis_data.get('detailed_findings', []), 1):
            result = finding.get('result', {})
            status_icon = "✅" if result.get('is_compliant') else "❌"
            
            elements.append(Paragraph(f"{i}. {status_icon} {finding['aspect']}", styles['Heading3']))
            elements.append(Paragraph(f"Status: {'COMPLIANT' if result.get('is_compliant') else 'NON-COMPLIANT'} | "
                                    f"Confidence: {result.get('confidence_score', 0)*100:.1f}%", styles['Normal']))
            
            if result.get('explanation'):
                elements.append(Paragraph(f"Analysis: {result['explanation']}", styles['Normal']))
            
            if result.get('document_evidence') and 'TIDAK DITEMUKAN' not in result['document_evidence']:
                elements.append(Paragraph(f"Evidence: \"{result['document_evidence'][:200]}...\"", styles['Normal']))
            
            elements.append(Spacer(1, 6))
        
        return elements

    def _create_pdf_recommendations(self, analysis_data: dict, heading_style, styles) -> list:
        """Create PDF recommendations elements"""
        elements = []
        
        elements.append(Paragraph('✅ ACTIONABLE RECOMMENDATIONS', heading_style))
        
        for i, rec in enumerate(analysis_data.get('recommendations', [])[:15], 1):
            if any(indicator in rec for indicator in ['🚨', '⚠️', '💡']):
                elements.append(Paragraph(rec, styles['Heading4']))
            else:
                elements.append(Paragraph(f"{i}. {rec}", styles['Normal']))
        
        elements.append(Spacer(1, 12))
        
        return elements

    def _create_pdf_appendix(self, analysis_data: dict, heading_style, styles) -> list:
        """Create PDF appendix elements"""
        elements = []
        
        elements.append(Paragraph('📚 APPENDIX', heading_style))
        
        # Metadata table
        metadata_data = [
            ['Standards Analyzed', ', '.join(analysis_data.get('analyzed_standards', []))],
            ['Analysis Method', 'AI Multi-Agent Adaptive Analysis'],
            ['System Version', 'ReguBot v2.0 Enhanced'],
            ['Total Aspects', str(len(analysis_data.get('detailed_findings', [])))],
            ['Processing Type', 'Offline - Secure']
        ]
        
        metadata_table = Table(metadata_data)
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.whitesmoke),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        
        elements.append(metadata_table)
        elements.append(Spacer(1, 12))
        
        # Footer
        elements.append(Paragraph('--- Enhanced Report by ReguBot AI Compliance Checker ---', styles['Normal']))
        
        return elements

    def enhanced_cover_page(self, doc, analysis_data: dict, session_id: str):
        """Enhanced cover page dengan informasi lebih lengkap"""
        # Main Title
        title = doc.add_heading('🛡️ LAPORAN AUDIT COMPLIANCE', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Inches(0.6)
        title_run.font.color.rgb = RGBColor(0x1f, 0x4e, 0x79)
        # Subtitle
        subtitle = doc.add_heading('ReguBot Enhanced Multi-Agent AI Analysis', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Document Analysis Summary
        doc_analysis = analysis_data.get('document_analysis', {})
        analysis_summary = self._create_analysis_summary(analysis_data)
        
        # Enhanced Info Table
        info_table = doc.add_table(rows=10, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        info_data = [
            ['📄 Dokumen Dianalisis', analysis_data.get('original_filename', 'document')],
            ['📊 Compliance Score', f"{analysis_data.get('compliance_score', 0)}%"],
            ['🔍 Aspek Dianalisis', f"{len(analysis_data.get('detailed_findings', []))} aspek compliance"],
            ['📈 Tingkat Compliance', analysis_summary['compliance_level']],
            ['🎯 Issues Ditemukan', f"{len(analysis_data.get('issues', []))} issues"],
            ['✅ Items Compliant', f"{len(analysis_data.get('compliant_items', []))} items"],
            ['📅 Tanggal Analisis', datetime.now().strftime('%d %B %Y')],
            ['⏰ Waktu Analisis', datetime.now().strftime('%H:%M:%S WIB')],
            ['🤖 AI Engine', 'ReguBot Multi-Agent v2.0 Enhanced'],
            ['🔒 Security Level', 'Offline Processing & Data Privacy Protected']
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)
            info_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Compliance Level Badge
        compliance_score = analysis_data.get('compliance_score', 0)
        if compliance_score >= 80:
            badge_text = "🟢 COMPLIANCE LEVEL: EXCELLENT"
        elif compliance_score >= 60:
            badge_text = "🟡 COMPLIANCE LEVEL: GOOD"
        elif compliance_score >= 40:
            badge_text = "🟠 COMPLIANCE LEVEL: NEEDS IMPROVEMENT"
        else:
            badge_text = "🔴 COMPLIANCE LEVEL: REQUIRES IMMEDIATE ATTENTION"
        
        badge_para = doc.add_paragraph(badge_text)
        badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        badge_run = badge_para.runs[0]
        badge_run.font.bold = True
        badge_run.font.size = Inches(0.25)
        
        doc.add_page_break()

    def _add_enhanced_cover_page(self, doc, analysis_data: dict, session_id: str):
        """Add a professional cover page to the DOCX report"""
        doc.add_heading('ReguBot Enhanced Compliance Audit Report', 0)
        doc.add_paragraph(f"Session ID: {session_id}", style='Intense Quote')
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style='Intense Quote')
        doc.add_paragraph(f"Document: {analysis_data.get('original_filename', 'Unknown')}", style='Intense Quote')
        doc.add_paragraph(f"Compliance Standards: {', '.join(analysis_data.get('analysis_metadata', {}).get('standards_analyzed', []))}", style='Intense Quote')
        doc.add_paragraph(f"System Version: {analysis_data.get('analysis_metadata', {}).get('system_version', 'ReguBot v2.0 Enhanced')}", style='Intense Quote')
        doc.add_paragraph("")
        doc.add_paragraph("This report was generated by ReguBot AI Multi-Agent System with Adaptive Analysis and Confidence-Weighted Scoring.")
        doc.add_paragraph("")
        doc.add_paragraph("Confidential - For authorized use only.")

    def _add_enhanced_executive_summary(self, doc, analysis_data: dict):
        """Enhanced executive summary dengan business insights"""
        doc.add_heading('📋 EXECUTIVE SUMMARY', level=1)
        
        analysis_summary = self._create_analysis_summary(analysis_data)
        doc_analysis = analysis_data.get('document_analysis', {})
        
        # Overall Assessment
        summary_text = f"""
{analysis_summary['status_icon']} OVERALL COMPLIANCE STATUS: {analysis_summary['compliance_level']}

Sistem ReguBot Enhanced AI telah melakukan analisis komprehensif terhadap dokumen yang disubmit menggunakan pendekatan adaptive multi-agent analysis. Dokumen diidentifikasi sebagai "{doc_analysis.get('document_type', 'General Document')}" dengan tingkat kompleksitas {doc_analysis.get('complexity_score', 0):.1f}/1.0.

🎯 KEY FINDINGS:
• Total aspek compliance yang dianalisis: {len(analysis_data.get('detailed_findings', []))} aspek (adaptive based on document content)
• Compliance score: {analysis_data.get('compliance_score', 0)}% (weighted scoring system)
• Issues yang memerlukan perhatian: {len(analysis_data.get('issues', []))} items
• Aspek yang sudah compliant: {len(analysis_data.get('compliant_items', []))} items
• Bahasa dokumen: {doc_analysis.get('language', 'Unknown')}
• Tema utama: {', '.join(doc_analysis.get('themes', [])[:3])}

📊 COMPLIANCE DISTRIBUTION:
        """
        
        # Add detailed breakdown
        if analysis_data.get('aspect_scores'):
            summary_text += "\n• Breakdown per aspek:\n"
            for aspect_key, score_info in analysis_data.get('aspect_scores', {}).items():
                aspect_name = next((finding['aspect'] for finding in analysis_data.get('detailed_findings', []) 
                                  if finding['aspect_key'] == aspect_key), aspect_key)
                summary_text += f"  - {aspect_name}: {score_info['score']*100:.1f}% (weight: {score_info['weight']:.2f})\n"
        
        summary_text += f"""
💡 BUSINESS IMPACT ASSESSMENT:
{analysis_summary['business_impact']}

🚀 IMMEDIATE ACTION REQUIRED:
{analysis_summary['immediate_actions']}

📈 EXPECTED IMPROVEMENT POTENTIAL:
Dengan implementasi rekomendasi yang diberikan, estimated compliance score dapat meningkat hingga {min(100, analysis_data.get('compliance_score', 0) + 30)}%.
        """
        
        doc.add_paragraph(summary_text.strip())
        doc.add_paragraph()

    def _add_document_analysis_overview(self, doc, analysis_data: dict):
        """Add document analysis overview section"""
        doc.add_heading('📄 DOCUMENT ANALYSIS OVERVIEW', level=1)
        
        doc_analysis = analysis_data.get('document_analysis', {})
        
        # Document characteristics table
        char_table = doc.add_table(rows=6, cols=2)
        char_table.style = 'Light List Accent 1'
        
        char_data = [
            ['Document Type', doc_analysis.get('document_type', 'Unknown')],
            ['Language Detected', doc_analysis.get('language', 'Unknown')],
            ['Word Count', f"{doc_analysis.get('word_count', 0):,} words"],
            ['Complexity Score', f"{doc_analysis.get('complexity_score', 0):.2f}/1.0"],
            ['Main Themes', ', '.join(doc_analysis.get('themes', [])[:5])],
            ['Sections Identified', f"{len(doc_analysis.get('sections', []))} sections"]
        ]
        
        for i, (label, value) in enumerate(char_data):
            char_table.cell(i, 0).text = label
            char_table.cell(i, 1).text = str(value)
            char_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        
        # Content themes analysis
        if doc_analysis.get('themes'):
            doc.add_paragraph()
            doc.add_heading('🎯 Content Themes Analysis', level=2)
            themes_text = f"""
Analisis AI mengidentifikasi {len(doc_analysis.get('themes', []))} tema utama dalam dokumen:

"""
            for i, theme in enumerate(doc_analysis.get('themes', [])[:8], 1):
                themes_text += f"{i}. {theme}\n"
            
            themes_text += f"""
Tema-tema ini digunakan untuk menentukan aspek compliance yang paling relevan untuk dianalisis, 
memastikan analisis yang focused dan practical sesuai dengan nature dokumen.
"""
            doc.add_paragraph(themes_text.strip())
        
        doc.add_paragraph()

    def _add_compliance_score_breakdown(self, doc, analysis_data: dict):
        """Add detailed compliance score breakdown"""
        doc.add_heading('📊 COMPLIANCE SCORE BREAKDOWN', level=1)
        
        # Main score display
        score_para = doc.add_paragraph(f"Overall Compliance Score: {analysis_data.get('compliance_score', 0)}%")
        score_run = score_para.runs[0]
        score_run.font.bold = True
        score_run.font.size = Inches(0.3)
        
        # Weighted scoring explanation
        doc.add_paragraph("""
🔍 SCORING METHODOLOGY:
Sistem menggunakan weighted scoring system dimana setiap aspek compliance memiliki bobot berbeda berdasarkan:
• Relevansi terhadap jenis dokumen
• Tingkat risiko compliance
• Frekuensi kemunculan dalam dokumen
• Impact terhadap data protection
        """)
        
        # Score breakdown table
        if analysis_data.get('aspect_scores'):
            breakdown_table = doc.add_table(rows=len(analysis_data['aspect_scores']) + 1, cols=4)
            breakdown_table.style = 'Medium Grid 1 Accent 1'
            
            # Headers
            headers = ['Aspek Compliance', 'Score (%)', 'Weight', 'Contribution']
            for i, header in enumerate(headers):
                cell = breakdown_table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data rows
            row_idx = 1
            for aspect_key, score_info in analysis_data['aspect_scores'].items():
                aspect_name = next((finding['aspect'] for finding in analysis_data.get('detailed_findings', []) 
                                  if finding['aspect_key'] == aspect_key), aspect_key.replace('_', ' ').title())
                
                breakdown_table.cell(row_idx, 0).text = aspect_name
                breakdown_table.cell(row_idx, 1).text = f"{score_info['score']*100:.1f}%"
                breakdown_table.cell(row_idx, 2).text = f"{score_info['weight']:.2f}"
                breakdown_table.cell(row_idx, 3).text = f"{score_info['weighted_contribution']*100:.1f}%"
                
                # Center align numeric columns
                for col in [1, 2, 3]:
                    breakdown_table.cell(row_idx, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                row_idx += 1
        
        doc.add_paragraph()

    def _add_enhanced_detailed_findings(self, doc, analysis_data: dict):
        """Enhanced detailed findings dengan evidence yang jelas"""
        doc.add_heading('🔍 DETAILED COMPLIANCE FINDINGS', level=1)
        
        detailed_findings = analysis_data.get('detailed_findings', [])
        
        for i, finding in enumerate(detailed_findings, 1):
            result = finding.get('result', {})
            
            # Aspect header dengan status
            status_icon = "✅" if result.get('is_compliant') else "❌"
            confidence = result.get('confidence_score', 0) * 100
            
            aspect_title = f"{i}. {status_icon} {finding['aspect']} (Confidence: {confidence:.1f}%)"
            doc.add_heading(aspect_title, level=2)
            
            # Compliance status dengan detail
            status_text = "COMPLIANT" if result.get('is_compliant') else "NON-COMPLIANT"
            severity = result.get('severity', 'MEDIUM')
            
            status_para = doc.add_paragraph(f"Status: {status_text} | Severity: {severity} | Weight: {finding.get('weight', 0):.2f}")
            status_para.runs[0].font.bold = True
            
            # Detailed explanation
            if result.get('explanation'):
                doc.add_paragraph(f"📝 Analysis: {result['explanation']}")
            
            # Document evidence dengan formatting yang lebih baik
            if result.get('document_evidence') and result['document_evidence'] not in ['TIDAK DITEMUKAN', 'TIDAK DAPAT DIANALISIS']:
                doc.add_paragraph("📄 Evidence from Document:")
                evidence_para = doc.add_paragraph(f'"{result["document_evidence"]}"')
                evidence_para.style = 'Intense Quote'
            
            # Found vs Missing elements
            if result.get('found_elements'):
                doc.add_paragraph(f"✅ Elements Found: {', '.join(result['found_elements'])}")
            
            if result.get('missing_elements'):
                doc.add_paragraph(f"❌ Missing Elements: {', '.join(result['missing_elements'])}")
            
            # Document excerpts
            excerpts = finding.get('document_excerpts', [])
            if excerpts:
                doc.add_paragraph(f"🔍 Relevant Document Excerpts ({len(excerpts)} found):")
                for j, excerpt in enumerate(excerpts[:2], 1):  # Show top 2 excerpts
                    excerpt_text = excerpt.get('text', '')[:200] + "..." if len(excerpt.get('text', '')) > 200 else excerpt.get('text', '')
                    doc.add_paragraph(f"{j}. \"{excerpt_text}\" (Score: {excerpt.get('score', 0):.2f})")
            
            # Standards applied
            standards_applied = finding.get('standards_applied', [])
            references = []
            details = []
            for std in standards_applied:
                ref = std.get('reference') or std.get('article') or std.get('source')
                if ref:
                    references.append(ref)
                # Add detail: title/section and content
                title = std.get('title', '')
                content = std.get('content', '')
                if title:
                    details.append(f"• {title}")
                if content:
                    details.append(f"  Bunyi: {content[:300]}{'...' if len(content)>300 else ''}")
            # Also add reference from compliance result if available
            result_ref = finding.get('result', {}).get('reference', '')
            if result_ref and result_ref not in references:
                references.append(result_ref)
            if references:
                doc.add_paragraph(f"📚 Standards Referenced: {', '.join(references)}")
            if details:
                doc.add_paragraph("📖 Detail Regulasi:")
                for detail in details:
                    doc.add_paragraph(detail)
            
            # Specific recommendations for this aspect
            if result.get('recommendations'):
                doc.add_paragraph("💡 Specific Recommendations:")
                for rec in result['recommendations'][:3]:  # Limit to 3 recommendations
                    doc.add_paragraph(f"• {rec}")
            
            doc.add_paragraph()

    def _add_risk_assessment_analysis(self, doc, analysis_data: dict):
        """Enhanced risk assessment dengan impact analysis"""
        doc.add_heading('⚠️ RISK ASSESSMENT & IMPACT ANALYSIS', level=1)
        
        issues = analysis_data.get('issues', [])
        
        # Risk categorization
        high_risk = [issue for issue in issues if issue.get('severity') == 'HIGH']
        medium_risk = [issue for issue in issues if issue.get('severity') == 'MEDIUM'] 
        low_risk = [issue for issue in issues if issue.get('severity') == 'LOW']
        
        # Risk summary
        doc.add_paragraph(f"""
🎯 RISK SUMMARY:
• High Risk Issues: {len(high_risk)} (Immediate action required)
• Medium Risk Issues: {len(medium_risk)} (Should be addressed within 30 days)
• Low Risk Issues: {len(low_risk)} (Can be addressed in next review cycle)
""")
        
        # Detailed risk analysis
        if high_risk:
            doc.add_heading('🚨 HIGH RISK ISSUES', level=2)
            for i, issue in enumerate(high_risk, 1):
                doc.add_paragraph(f"""
{i}. {issue['aspect']}
   Risk: {issue.get('explanation', 'N/A')}
   Business Impact: Potential regulatory fines, legal liability, user trust damage
   Recommended Timeline: Immediate (within 7 days)
""")
        
        if medium_risk:
            doc.add_heading('⚠️ MEDIUM RISK ISSUES', level=2)
            for i, issue in enumerate(medium_risk, 1):
                doc.add_paragraph(f"""
{i}. {issue['aspect']}
   Risk: {issue.get('explanation', 'N/A')}
   Business Impact: Compliance gaps, potential user complaints
   Recommended Timeline: 30 days
""")
        
        # Risk mitigation strategies
        doc.add_heading('🛡️ RISK MITIGATION STRATEGIES', level=2)
        doc.add_paragraph("""
1. IMMEDIATE ACTIONS (0-7 days):
   • Address all HIGH risk issues
   • Implement temporary controls if needed
   • Notify relevant stakeholders

2. SHORT-TERM ACTIONS (1-4 weeks):
   • Develop comprehensive remediation plan
   • Address MEDIUM risk issues
   • Update internal procedures

3. LONG-TERM ACTIONS (1-3 months):
   • Implement permanent controls
   • Conduct compliance training
   • Establish monitoring mechanisms
""")
        
        doc.add_paragraph()

    def _add_prioritized_recommendations(self, doc, analysis_data: dict):
        """Enhanced recommendations dengan prioritization yang jelas"""
        doc.add_heading('✅ PRIORITIZED ACTION PLAN', level=1)
        
        recommendations = analysis_data.get('recommendations', [])
        
        # Categorize recommendations
        priority_sections = []
        current_section = None
        current_items = []
        
        for rec in recommendations:
            if any(indicator in rec for indicator in ['🚨', '⚠️', '💡', '🇪🇺', '🇮🇩', '🛡️', '🔧', '✨']):
                if current_section:
                    priority_sections.append({'title': current_section, 'items': current_items})
                current_section = rec
                current_items = []
            else:
                current_items.append(rec)
        
        if current_section:
            priority_sections.append({'title': current_section, 'items': current_items})
        
        # Display organized recommendations
        for section in priority_sections:
            if section['title'] and section['items']:
                doc.add_heading(section['title'], level=2)
                for item in section['items']:
                    doc.add_paragraph(item)
                doc.add_paragraph()
        
        # Implementation timeline
        doc.add_heading('📅 IMPLEMENTATION TIMELINE', level=2)
        timeline_table = doc.add_table(rows=4, cols=3)
        timeline_table.style = 'Light Grid Accent 1'
        
        timeline_data = [
            ['Phase', 'Timeframe', 'Focus Areas'],
            ['IMMEDIATE', '0-7 days', 'High risk issues, critical gaps'],
            ['SHORT-TERM', '1-4 weeks', 'Medium risk issues, process improvements'],
            ['LONG-TERM', '1-3 months', 'Optimization, monitoring, training']
        ]
        
        for i, (phase, timeframe, focus) in enumerate(timeline_data):
            timeline_table.cell(i, 0).text = phase
            timeline_table.cell(i, 1).text = timeframe
            timeline_table.cell(i, 2).text = focus
            if i == 0:  # Header row
                for j in range(3):
                    timeline_table.cell(i, j).paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()

    def _add_implementation_roadmap(self, doc, analysis_data: dict):
        """Add implementation roadmap section"""
        doc.add_heading('🗺️ IMPLEMENTATION ROADMAP', level=1)
        
        issues = analysis_data.get('issues', [])
        compliance_score = analysis_data.get('compliance_score', 0)
        
        # Roadmap based on current compliance level
        if compliance_score < 40:
            roadmap_type = "COMPREHENSIVE OVERHAUL"
            duration = "3-6 months"
            phases = [
                "Phase 1 (Month 1): Critical compliance gaps",
                "Phase 2 (Month 2-3): Core policy development", 
                "Phase 3 (Month 4-5): Implementation & training",
                "Phase 4 (Month 6): Monitoring & optimization"
            ]
        elif compliance_score < 70:
            roadmap_type = "TARGETED IMPROVEMENTS"
            duration = "2-3 months"
            phases = [
                "Phase 1 (Week 1-2): High priority issues",
                "Phase 2 (Month 1): Policy enhancements",
                "Phase 3 (Month 2-3): Process improvements"
            ]
        else:
            roadmap_type = "OPTIMIZATION & MAINTENANCE"
            duration = "1-2 months"
            phases = [
                "Phase 1 (Week 1-2): Minor adjustments",
                "Phase 2 (Month 1): Best practice implementation", 
                "Phase 3 (Month 2): Continuous monitoring setup"
            ]
        
        doc.add_paragraph(f"""
🎯 RECOMMENDED APPROACH: {roadmap_type}
⏱️ ESTIMATED DURATION: {duration}
📊 CURRENT COMPLIANCE LEVEL: {compliance_score}%
🎯 TARGET COMPLIANCE LEVEL: {min(100, compliance_score + 40)}%

📋 IMPLEMENTATION PHASES:
""")
        
        for phase in phases:
            doc.add_paragraph(f"• {phase}")
        
        # Success metrics
        doc.add_paragraph(f"""

📈 SUCCESS METRICS:
• Compliance score improvement: Target {min(100, compliance_score + 40)}%
• High-risk issues resolved: 100%
• Medium-risk issues resolved: 80%
• Documentation completeness: 95%
• Stakeholder awareness: 90%
• Process automation: 70%
""")
        
        doc.add_paragraph()
