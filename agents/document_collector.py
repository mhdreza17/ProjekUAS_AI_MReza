import os
import fitz  # PyMuPDF
from docx import Document
import pytesseract
from PIL import Image
import io
from .base_agent import BaseAgent

class DocumentCollectorAgent(BaseAgent):
    """Agent untuk mengumpulkan dan memproses dokumen"""
    
    def __init__(self):
        super().__init__("DocumentCollector")
        self.supported_formats = ['.pdf', '.docx', '.txt']
        self.upload_folder = 'uploads'  # Tambah untuk rekonstruksi path
        
    def process(self, filepath: str):
        """
        Ekstrak teks dari dokumen yang diupload
        
        Prompt Internal: Anda adalah Document Collector Agent. Tugas Anda: 
        terima dokumen dari pengguna, ekstrak seluruh teksnya secara akurat, 
        dan kirim ke Compliance Checker Agent. Gunakan OCR hanya jika dokumen 
        adalah gambar. Pastikan tidak ada bagian yang terpotong atau terlewat. 
        Output: teks lengkap dalam format UTF-8, siap dianalisis.
        """
        self.set_status("processing")
        self.log_action("Starting document extraction", filepath)
        
        try:
            # Fix: Jika filepath tidak exist (kemungkinan hanya session_id), rekonstruksi full path
            if not os.path.exists(filepath):
                self.log_action("Reconstructing filepath", "Input seems to be session_id only")
                files = [f for f in os.listdir(self.upload_folder) if f.startswith(filepath + '_')]
                if not files:
                    raise ValueError(f"File not found for session: {filepath}")
                filepath = os.path.join(self.upload_folder, files[0])  # Ambil file pertama yang match
                self.log_action("Reconstructed filepath", filepath)
            
            # Sekarang ekstensi dari filepath yang benar
            file_ext = os.path.splitext(filepath)[1].strip().lower()
            self.log_action("Detected file extension", file_ext)
            self.log_action("File path received", filepath)
            
            if file_ext in self.supported_formats:
                if file_ext == '.pdf':
                    text = self._extract_from_pdf(filepath)
                elif file_ext == '.docx':
                    text = self._extract_from_docx(filepath)
                elif file_ext == '.txt':
                    text = self._extract_from_txt(filepath)
            else:
                raise ValueError(f"Format file tidak didukung: {file_ext} (filepath: {filepath})")
            
            # Validasi hasil ekstraksi
            if not text or len(text.strip()) < 10:
                self.log_action("Warning", "Teks yang diekstrak sangat pendek atau kosong")
                
            self.set_status("completed")
            self.log_action("Document extraction completed", f"Extracted {len(text)} characters")
            
            return {
                'success': True,
                'text': text,
                'file_type': file_ext,
                'char_count': len(text),
                'word_count': len(text.split())
            }
            
        except Exception as e:
            self.set_status("error")
            self.log_action("Error during extraction", str(e))
            return {
                'success': False,
                'error': str(e)
            }
    
    def _extract_from_pdf(self, filepath: str) -> str:
        """Ekstrak teks dari file PDF"""
        text = ""
        try:
            doc = fitz.open(filepath)
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text()
                
                # Jika halaman kosong atau sangat sedikit teks, coba OCR
                if len(page_text.strip()) < 50:
                    self.log_action("OCR Processing", f"Page {page_num + 1} - low text content")
                    page_text = self._ocr_page(page)
                
                text += page_text + "\n"
            
            doc.close()
            return text.strip()
            
        except Exception as e:
            self.log_action("PDF extraction error", str(e))
            raise e
    
    def _extract_from_docx(self, filepath: str) -> str:
        """Ekstrak teks dari file DOCX"""
        try:
            doc = Document(filepath)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Ekstrak teks dari tabel jika ada
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            self.log_action("DOCX extraction error", str(e))
            raise e
    
    def _extract_from_txt(self, filepath: str) -> str:
        """Ekstrak teks dari file TXT"""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Coba encoding lain jika UTF-8 gagal
            with open(filepath, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _ocr_page(self, page) -> str:
        """Lakukan OCR pada halaman PDF yang berupa gambar"""
        try:
            # Convert PDF page to image
            pix = page.get_pixmap()
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            # Perform OCR
            text = pytesseract.image_to_string(img, lang='ind+eng')
            return text
            
        except Exception as e:
            self.log_action("OCR error", str(e))
            return ""